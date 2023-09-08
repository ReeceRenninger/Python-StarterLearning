# python doc x contains multiple data types
# https://python-docx.readthedocs.io/en/latest/user/quickstart.html

import docx
import os

os.chdir('/home/reece/PracticeFiles')

document = docx.Document('demo.docx') # create document object

print(document.paragraphs[0].text) # returns text of paragraph at index 0

print(document.paragraphs[1].text) # returns text of paragraph at index 1

paragraph = document.paragraphs[1] # create paragraph object that contains runs

#! a run is a group of characters with the same style, so if you bold a word, it will be a new run, then if you go back to normal text, it will be a new run
print(paragraph.runs[1].text) # returns text of the 1st run
print(paragraph.runs[2].text) # returns text of the 2nd run
print(paragraph.runs[1].bold) # returns boolean if run is bold
print(paragraph.runs[2].bold) # returns boolean if run is bold

paragraph.runs[3].underline = True # underline the 3rd run
paragraph.runs[3].text = 'italic and underlined' # change text of 3rd run

document.save('/home/reece/PracticeFiles/demo2.docx') # save changes to new file

# create new document
newDocument = docx.Document()
newDocument.add_paragraph('Hello this is a paragraph')
newDocument.add_paragraph('This is another paragraph')
newDocument.save('/home/reece/PracticeFiles/demo3.docx')


# add run
p = newDocument.paragraphs[0]
p.add_run('This is a new run')
p.runs[1].bold = True
newDocument.save('/home/reece/PracticeFiles/demo4.docx')

# modify existing document
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for paragraphs in doc.paragraphs:
        fullText.append(paragraphs.text) # append each paragraph to fullText list
    return '\n'.join(fullText) # join each paragraph with a new line

print(getText('/home/reece/PracticeFiles/demo.docx')) # print all text in demo.docx